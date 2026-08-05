"""md-to-docx - convert a markdown file into a formatted Word document.

Built for Mick's PAIDA system (Cedric). Turns any vault / project markdown file
into a clean, readable .docx with real Word structure: heading styles, native
tables, embedded images, bullet and numbered lists, hyperlinks, code blocks,
the brand logo in the header and a right-aligned "Page x of y" footer.

Optionally exports a matching PDF via Word (pywin32 COM).

Usage:
    python md_to_docx.py "<input.md>"
    python md_to_docx.py "<input.md>" --pdf
    python md_to_docx.py "<input.md>" --brand ai            (diy-investors.ai)
    python md_to_docx.py "<input.md>" --brand none          (no header logo)
    python md_to_docx.py "<input.md>" --final               (the copy Mick sends)
    python md_to_docx.py "<input.md>" -o "<output.docx>" --landscape
    python md_to_docx.py "<input.md>" --no-footer           (no footer at all)
    python md_to_docx.py "<input.md>" --frontmatter         (keep YAML as a block)

Branding (added 2026-08-05):
    --brand com   diy-investors.com logo (DEFAULT - Inner Circle, Plaza Group)
    --brand ai    diy-investors.ai logo (AI for Investing material)
    --brand none  no logo
The logo sits top LEFT of the header, at 25 percent of the usable page width,
original aspect ratio preserved.

Footer - draft vs final (Mick's rule, 2026-08-05):
    Every document gets "Page x of y", RIGHT aligned, using live Word PAGE and
    NUMPAGES fields so it stays correct after editing and in the PDF.

    DRAFT (the default) ALSO carries the file path on the left, labelled:
        DRAFT - C:\\...\\file.docx - Created: YYYY.MM.DD
    Mick works through several iterations of most documents and uses that path
    to find the file again in Obsidian, so it is on by default.

    FINAL (--final) drops the path entirely - just the logo and the page
    numbers. Use it for the copy that goes to members: they must never see
    Mick's local vault paths. The DRAFT label is what stops a working copy
    passing as a finished one, which a bare path line never did.

    A PDF exported with --pdf is produced FROM the DOCX, so it inherits
    whichever mode was used. There is no separate PDF rule to remember.

Markdown supported:
    YAML frontmatter (skipped by default)
    # .. ###### headings
    paragraphs with **bold**, *italic*, ***both***, `code`
    [links](https://...) as real Word hyperlinks
    - and * bullets (two levels), 1. numbered lists
    | pipe | tables | with header separator row
    ![alt](image.png) images, resolved relative to the markdown file
    > blockquotes
    ``` fenced code blocks
    --- horizontal rules
    <!-- pagebreak --> forces a page break

ASCII policy: typographic punctuation (em dash, curly quotes, ellipsis) is
converted to ASCII automatically. The pound and cent signs are left alone -
Mick's standing decision that they are legitimate content, matching the
REVIEW_IGNORE list in non-ascii-sweep. Anything else non-ASCII is REPORTED with
its line number and left untouched, so the source can be fixed at source.

Version: 1.3 (2026-08-05 - brand header logo, Page x of y footer, draft/final)
"""

import argparse
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

# ---------------------------------------------------------------- constants --

HEADING_COLOUR = RGBColor(0x1F, 0x38, 0x64)   # dark blue, matches house style
LINK_COLOUR = "0563C1"
HEADER_SHADE = "DCE6F1"
GREY = RGBColor(0x80, 0x80, 0x80)

HEADING_SIZES = {1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(11), 5: Pt(11), 6: Pt(11)}

INLINE_RE = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`)")
LINK_RE = re.compile(r"(?<!!)\[([^\]]+)\]\(([^)\s]+)\)")
IMG_RE = re.compile(r"^!\[(.*?)\]\((.+?)\)\s*$")
# Obsidian embed: ![[image.png]] or ![[image.png|caption]] - resolved by
# filename anywhere in the vault, the way Obsidian itself resolves it.
WIKI_IMG_RE = re.compile(r"^!\[\[([^\]|]+?)(?:\|([^\]]*))?\]\]\s*$")
IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".emf",
              ".wmf")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^(\s*)[-*+]\s+(.*)$")
NUMBER_RE = re.compile(r"^(\s*)\d+[.)]\s+(.*)$")
HRULE_RE = re.compile(r"^\s*([-*_])(\s*\1){2,}\s*$")
PAGEBREAK_RE = re.compile(r"^\s*<!--\s*pagebreak\s*-->\s*$", re.I)

# Typographic characters that are ALWAYS converted (per the ASCII-only rule).
# Written as escape sequences so THIS file stays pure ASCII and never trips the
# weekly non-ascii-sweep over the vault.
ASCII_MAP = {
    "\u2014": "-",    # em dash
    "\u2013": "-",    # en dash
    "\u2212": "-",    # minus sign
    "\u2010": "-",    # hyphen
    "\u2011": "-",    # non-breaking hyphen
    "\u2018": "'",    # left single quote
    "\u2019": "'",    # right single quote / apostrophe
    "\u201a": "'",    # single low quote
    "\u201b": "'",    # single high-reversed quote
    "\u201c": '"',    # left double quote
    "\u201d": '"',    # right double quote
    "\u201e": '"',    # double low quote
    "\u2026": "...",  # ellipsis
    "\u00a0": " ",    # non-breaking space
    "\u2007": " ",    # figure space
    "\u202f": " ",    # narrow no-break space
    "\u2022": "-",    # bullet
    "\u00b7": "-",    # middle dot
    "\u2043": "-",    # hyphen bullet
    "\u00ad": "",     # soft hyphen
    "\u200b": "",     # zero-width space
    "\ufeff": "",     # byte-order mark
    "\u2032": "'",    # prime
    "\u2033": '"',    # double prime
    "\u00b4": "'",    # acute accent
    "\u02bc": "'",    # modifier apostrophe
}

# Non-ASCII that is legitimate content and must NOT be flagged or altered.
# Matches the REVIEW_IGNORE decision in non-ascii-sweep (2026-08-02).
ASCII_ALLOW = {"\u00a3", "\u00a2"}   # pound sign, cent sign

# ------------------------------------------------------------------ branding --
#
# Which logo goes on which document (Mick's rule, 2026-08-05):
#   com -> diy-investors.com  : Inner Circle and Plaza Group material
#   ai  -> diy-investors.ai   : AI for Investing material
#
# Logos are shipped inside this skill (assets/) so the skill is self-contained
# and does not break if the Documents project folders are reorganised. The
# original masters in "0 - AI Logos n Podcast Covers" are kept as a fallback.
HERE = os.path.dirname(os.path.abspath(__file__))
LOGO_MASTERS = os.path.join(
    os.path.expanduser("~"), "Documents", "0.1 - Projects (n)",
    "0 - AI Logos n Podcast Covers", "0 - Logos")

BRAND_LOGOS = {
    "com": [
        os.path.join(HERE, "assets", "logo-diy-investors-com.jpg"),
        os.path.join(LOGO_MASTERS, "DIY-Logo_290 x 58px_for Covers_White_JPG.jpg"),
    ],
    "ai": [
        os.path.join(HERE, "assets", "logo-diy-investors-ai.jpg"),
        os.path.join(LOGO_MASTERS,
                     "2024.12.18 - 800x200px_DIY.ai_v.02_Logo_white_JPG.jpg"),
    ],
}

# Logo width as a fraction of the usable page width (page minus both margins).
LOGO_WIDTH_FRACTION = 0.25


# ------------------------------------------------------------------ helpers --

def clean_ascii(text):
    """Convert typographic punctuation; report anything else non-ASCII."""
    for bad, good in ASCII_MAP.items():
        text = text.replace(bad, good)
    problems = []
    for lineno, line in enumerate(text.split("\n"), start=1):
        for ch in line:
            if ord(ch) > 127 and ch not in ASCII_ALLOW:
                problems.append((lineno, ch, hex(ord(ch))))
    return text, problems


def add_hyperlink(par, url, text, size=None):
    part = par.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), LINK_COLOUR)
    rpr.append(colour)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size.pt * 2)))
        rpr.append(sz)
    run.append(rpr)
    node = OxmlElement("w:t")
    node.text = text
    node.set(qn("xml:space"), "preserve")
    run.append(node)
    link.append(run)
    par._p.append(link)


def add_runs(par, text, cfg, bold=False, italic=False, size=None):
    """Add text to a paragraph honouring inline emphasis, code and links."""
    size = size or cfg["size"]
    pos = 0
    for m in LINK_RE.finditer(text):
        if m.start() > pos:
            _emphasis_runs(par, text[pos:m.start()], cfg, bold, italic, size)
        add_hyperlink(par, m.group(2), m.group(1), size)
        pos = m.end()
    if pos < len(text):
        _emphasis_runs(par, text[pos:], cfg, bold, italic, size)
    return par


def _emphasis_runs(par, text, cfg, bold, italic, size):
    for chunk in INLINE_RE.split(text):
        if not chunk:
            continue
        b, i, mono = bold, italic, False
        if chunk.startswith("***") and chunk.endswith("***") and len(chunk) > 6:
            chunk, b, i = chunk[3:-3], True, True
        elif chunk.startswith("**") and chunk.endswith("**") and len(chunk) > 4:
            chunk, b = chunk[2:-2], True
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            chunk, i = chunk[1:-1], True
        elif chunk.startswith("`") and chunk.endswith("`") and len(chunk) > 2:
            chunk, mono = chunk[1:-1], True
        run = par.add_run(chunk)
        run.bold = b
        run.italic = i
        run.font.name = "Consolas" if mono else cfg["font"]
        run.font.size = Pt(size.pt - 0.5) if mono else size


def add_hrule(doc):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(6)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pbdr.append(bottom)
    par._p.get_or_add_pPr().append(pbdr)


def shade_cell(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_sep_row(line):
    line = line.strip()
    return (line.startswith("|") or "|" in line) and \
        bool(re.match(r"^\|?[\s:|-]+\|[\s:|-]*$", line)) and "-" in line


def add_table(doc, rows, cfg):
    width = max(len(r) for r in rows)
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for c in range(width):
        cell = table.cell(0, c)
        cell.text = ""
        par = cell.paragraphs[0]
        par.paragraph_format.space_after = Pt(2)
        add_runs(par, header[c] if c < len(header) else "", cfg,
                 bold=True, size=cfg["table_size"])
        shade_cell(cell, HEADER_SHADE)
    for r, row in enumerate(body, start=1):
        for c in range(width):
            cell = table.cell(r, c)
            cell.text = ""
            par = cell.paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            add_runs(par, row[c] if c < len(row) else "", cfg,
                     size=cfg["table_size"])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def find_vault_root(start_dir):
    """Walk up from start_dir looking for the folder that holds .obsidian."""
    path = os.path.abspath(start_dir)
    while True:
        if os.path.isdir(os.path.join(path, ".obsidian")):
            return path
        parent = os.path.dirname(path)
        if parent == path:
            return None
        path = parent


_VAULT_INDEX = {}


def vault_lookup(name, src_dir):
    """Find a file by BASENAME anywhere in the vault, as Obsidian does.

    The index is built once per vault per run and skips the usual noise
    directories, so a big vault costs one walk, not one per embed.
    """
    root = find_vault_root(src_dir)
    if root is None:
        return None
    if root not in _VAULT_INDEX:
        index = {}
        skip = {".git", ".obsidian", ".trash", "__pycache__", "node_modules",
                ".venv"}
        for dirpath, dirnames, filenames in os.walk(root):
            dirnames[:] = [d for d in dirnames if d not in skip]
            for fn in filenames:
                index.setdefault(fn.lower(), os.path.join(dirpath, fn))
        _VAULT_INDEX[root] = index
    return _VAULT_INDEX[root].get(os.path.basename(name).lower())


def resolve_asset(path, src_dir, wiki=False):
    """Resolve an image reference to a file on disk, or None.

    Order: absolute path, then relative to the markdown file, then (for
    Obsidian ![[embeds]]) a vault-wide search on the filename.
    """
    if os.path.isabs(path) and os.path.exists(path):
        return path
    local = os.path.join(src_dir, path)
    if os.path.exists(local):
        return local
    if wiki:
        return vault_lookup(path, src_dir)
    return None


def usable_width(section):
    """Width of the printable area: page width less both margins (EMU)."""
    return section.page_width - section.left_margin - section.right_margin


def resolve_logo(brand):
    """Return the first logo file that exists for this brand, else None."""
    for path in BRAND_LOGOS.get(brand, []):
        if os.path.exists(path):
            return path
    return None


def add_brand_header(section, brand, cfg):
    """Put the brand logo top LEFT of the header, at 25% of the usable width.

    python-docx scales the height automatically when only a width is given, so
    the original aspect ratio is always preserved.
    """
    if brand in (None, "none"):
        return None
    logo = resolve_logo(brand)
    if logo is None:
        print("WARNING: no logo file found for brand '%s' - header left blank."
              % brand)
        return None
    header = section.header
    header.is_linked_to_previous = False
    par = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    par.text = ""
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.space_after = Pt(6)
    width = Emu(int(usable_width(section) * LOGO_WIDTH_FRACTION))
    par.add_run().add_picture(logo, width=width)
    return logo


def add_field(par, instruction):
    """Insert a live Word field (PAGE, NUMPAGES, ...) into a paragraph.

    Word evaluates these at layout time, so the numbers stay correct after
    editing and come through correctly in the exported PDF.
    """
    run = par.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, end):
        run._r.append(node)
    return run


def add_footer(doc, section, cfg, draft_text=None):
    """Footer: "Page x of y", hard RIGHT, on every document.

    draft_text is the DRAFT path line. It goes on its own line above the page
    numbers, because a full vault path plus page numbers does not fit on one A4
    line. Passed on drafts (the default) and omitted on --final copies.

    Note: Word's built-in Footer style carries its own centre and right tab
    stops, sized for US Letter with 1in margins. Those are cleared here and
    replaced with a right tab at the real page width, otherwise anything
    tabbed in a footer lands mid-page instead of at the right margin.
    """
    footer = section.footer
    footer.is_linked_to_previous = False
    par = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    par.text = ""
    par.paragraph_format.space_before = Pt(2)

    style_tabs = doc.styles["Footer"].paragraph_format.tab_stops
    style_tabs.clear_all()
    style_tabs.add_tab_stop(Emu(int(usable_width(section))),
                            WD_TAB_ALIGNMENT.RIGHT)

    if draft_text:
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        par.paragraph_format.space_after = Pt(0)
        run = par.add_run(draft_text)
        run.font.size = Pt(7.5)
        run.font.name = cfg["font"]
        run.font.color.rgb = GREY
        par = footer.add_paragraph()

    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    par.add_run("Page ")
    add_field(par, " PAGE ")
    par.add_run(" of ")
    add_field(par, " NUMPAGES ")
    for run in par.runs:
        if run.font.size is None:
            run.font.size = Pt(8.5)
            run.font.name = cfg["font"]
            run.font.color.rgb = GREY


def export_pdf(docx_path):
    """Export the DOCX to PDF.

    Delegates to docx_to_pdf.py, the skill's shared Word COM engine, so there is
    only one implementation of the export to maintain. That engine uses
    ExportAsFixedFormat, which carries the Word heading structure through as PDF
    bookmarks. Returns the pdf path, or None if the export could not run.
    """
    here = os.path.dirname(os.path.abspath(__file__))
    if here not in sys.path:
        sys.path.insert(0, here)
    try:
        from docx_to_pdf import convert_one
    except ImportError:
        print("WARNING: docx_to_pdf.py not found alongside this script "
              "- PDF skipped. The DOCX has still been saved.")
        return None
    return convert_one(docx_path)


# -------------------------------------------------------------- conversion --

def convert(src, out, args):
    src_dir = os.path.dirname(os.path.abspath(src))
    with open(src, "r", encoding="utf-8") as fh:
        raw = fh.read()

    raw, problems = clean_ascii(raw)
    if problems:
        print("NON-ASCII CHARACTERS FOUND (left untouched - fix at source):")
        seen = set()
        for lineno, ch, code in problems:
            key = (ch, lineno)
            if key in seen:
                continue
            seen.add(key)
            print("  line %d: %r (%s)" % (lineno, ch, code))

    lines = raw.replace("\r\n", "\n").split("\n")

    # Frontmatter
    frontmatter = []
    if lines and lines[0].strip() == "---":
        end = None
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                end = i
                break
        if end is not None:
            frontmatter = lines[1:end]
            lines = lines[end + 1:]

    cfg = {
        "font": args.font,
        "size": Pt(args.size),
        "table_size": Pt(args.table_size),
    }

    doc = Document()
    section = doc.sections[0]
    if args.landscape:
        section.page_width, section.page_height = Inches(11.69), Inches(8.27)
        section.orientation = 1
    else:
        section.page_width, section.page_height = Inches(8.27), Inches(11.69)
    for side in ("top", "bottom", "left", "right"):
        setattr(section, "%s_margin" % side, Inches(args.margin))

    logo = add_brand_header(section, args.brand, cfg)
    if logo:
        print("Header logo (%s): %s" % (args.brand, os.path.basename(logo)))

    normal = doc.styles["Normal"]
    normal.font.name = cfg["font"]
    normal.font.size = cfg["size"]
    normal.paragraph_format.space_after = Pt(6)

    if args.frontmatter and frontmatter:
        par = doc.add_paragraph()
        for run in add_runs(par, " | ".join(f.strip() for f in frontmatter
                                            if f.strip() and not f.startswith("- ")),
                            cfg, italic=True, size=Pt(8)).runs:
            run.font.color.rgb = GREY
        add_hrule(doc)

    i = 0
    title_done = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if PAGEBREAK_RE.match(stripped):
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue

        if HRULE_RE.match(stripped):
            add_hrule(doc)
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            j = i + 1
            code = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                code.append(lines[j])
                j += 1
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.25)
            par.paragraph_format.space_after = Pt(8)
            run = par.add_run("\n".join(code))
            run.font.name = "Consolas"
            run.font.size = Pt(cfg["size"].pt - 1.5)
            i = j + 1
            continue

        # Image - standard markdown ![alt](path) or Obsidian embed ![[file.png]]
        m = IMG_RE.match(stripped)
        wm = WIKI_IMG_RE.match(stripped) if not m else None
        if m or wm:
            wiki = wm is not None
            path = wm.group(1).strip() if wiki else m.group(2)
            if wiki and not path.lower().endswith(IMAGE_EXTS):
                # ![[Some Note]] - an embedded note, not an image. Not something
                # this converter can inline; leave a visible marker.
                par = doc.add_paragraph()
                add_runs(par, "[embedded note: %s]" % path, cfg)
                print("WARNING: embedded note skipped (not an image) -> %s" % path)
                i += 1
                continue
            full = resolve_asset(path, src_dir, wiki=wiki)
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_after = Pt(4)
            if full:
                par.add_run().add_picture(full, width=Inches(args.image_width))
            else:
                add_runs(par, "[image not found: %s]" % path, cfg)
                print("WARNING: image not found -> %s" % path)
            i += 1
            continue

        # Table
        if "|" in stripped and i + 1 < len(lines) and is_sep_row(lines[i + 1]):
            rows = [split_row(stripped)]
            j = i + 2
            while j < len(lines) and "|" in lines[j] and lines[j].strip():
                rows.append(split_row(lines[j]))
                j += 1
            add_table(doc, rows, cfg)
            i = j
            continue

        # Heading
        hm = HEADING_RE.match(stripped)
        if hm:
            level = len(hm.group(1))
            par = doc.add_paragraph(style="Heading %d" % min(level, 4))
            par.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
            par.paragraph_format.space_after = Pt(4)
            par.paragraph_format.keep_with_next = True
            for run in add_runs(par, hm.group(2), cfg, bold=True,
                                size=HEADING_SIZES[level]).runs:
                run.font.color.rgb = HEADING_COLOUR
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            text = [stripped.lstrip("> ").strip()]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith(">"):
                text.append(lines[j].strip().lstrip("> ").strip())
                j += 1
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.3)
            add_runs(par, " ".join(text), cfg, italic=True)
            i = j
            continue

        # Bullet / numbered list item (with wrapped continuation lines)
        bm = BULLET_RE.match(line)
        nm = NUMBER_RE.match(line) if not bm else None
        if bm or nm:
            match = bm or nm
            indent = len(match.group(1))
            text = match.group(2)
            j = i + 1
            while j < len(lines):
                nxt = lines[j]
                if not nxt.strip():
                    break
                if BULLET_RE.match(nxt) or NUMBER_RE.match(nxt) or \
                        HEADING_RE.match(nxt.strip()) or "|" in nxt.strip()[:2]:
                    break
                if not nxt.startswith(" "):
                    break
                text += " " + nxt.strip()
                j += 1
            if bm:
                style = "List Bullet 2" if indent >= 2 else "List Bullet"
            else:
                style = "List Number 2" if indent >= 2 else "List Number"
            par = doc.add_paragraph(style=style)
            par.paragraph_format.space_after = Pt(3)
            add_runs(par, text, cfg)
            i = j
            continue

        # Document title line: "**Header:** ..." (Ron / Nina report convention)
        if not title_done and stripped.startswith("**Header:**"):
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_after = Pt(10)
            for run in add_runs(par, stripped.replace("**Header:**", "").strip(),
                                cfg, bold=True, size=Pt(15)).runs:
                run.font.color.rgb = HEADING_COLOUR
            add_hrule(doc)
            title_done = True
            i += 1
            continue

        # Plain paragraph, joining wrapped lines
        text = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if not nxt or HEADING_RE.match(nxt) or BULLET_RE.match(lines[j]) or \
                    NUMBER_RE.match(lines[j]) or HRULE_RE.match(nxt) or \
                    nxt.startswith(("|", ">", "```", "![")):
                break
            text.append(nxt)
            j += 1
        par = doc.add_paragraph()
        if args.justify:
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(par, " ".join(text), cfg)
        i = j

    if not args.no_footer:
        draft_text = None
        if not args.final:
            draft_text = "DRAFT - %s - Created: %s" % (os.path.abspath(out),
                                                       args.date)
        add_footer(doc, section, cfg, draft_text)
        print("Footer: %s" % ("FINAL - page numbers only, no path"
                              if args.final else
                              "DRAFT - labelled path plus page numbers"))

    doc.save(out)
    return out


def main():
    ap = argparse.ArgumentParser(description="Convert markdown to a Word document.")
    ap.add_argument("source", help="path to the .md file")
    ap.add_argument("-o", "--output", help="output .docx path "
                                           "(default: same folder, same name)")
    ap.add_argument("--pdf", action="store_true", help="also export a PDF via Word")
    ap.add_argument("--brand", choices=("com", "ai", "none"), default="com",
                    help="header logo: com = diy-investors.com (Inner Circle / "
                         "Plaza Group, the default), ai = diy-investors.ai, "
                         "none = no logo")
    ap.add_argument("--final", action="store_true",
                    help="FINAL copy: page numbers only, no path in the footer. "
                         "Use for anything going to members. Without it the "
                         "document is a DRAFT and carries the labelled file "
                         "path so Mick can find it again in Obsidian")
    ap.add_argument("--no-footer", action="store_true",
                    help="omit the footer entirely (going-to-press copy)")
    ap.add_argument("--frontmatter", action="store_true",
                    help="render the YAML frontmatter as a grey metadata line")
    ap.add_argument("--landscape", action="store_true", help="landscape A4")
    ap.add_argument("--font", default="Calibri")
    ap.add_argument("--size", type=float, default=11.0, help="body font size (pt)")
    ap.add_argument("--table-size", type=float, default=8.5,
                    help="table font size (pt)")
    ap.add_argument("--image-width", type=float, default=6.6,
                    help="embedded image width (inches)")
    ap.add_argument("--margin", type=float, default=0.7, help="page margin (inches)")
    ap.add_argument("--no-justify", dest="justify", action="store_false",
                    help="left-align body text instead of justifying")
    ap.add_argument("--date", help="creation date shown in the DRAFT footer "
                                   "(YYYY.MM.DD); defaults to today")
    args = ap.parse_args()

    if not os.path.exists(args.source):
        sys.exit("ERROR: source not found - %s" % args.source)

    out = args.output or os.path.splitext(os.path.abspath(args.source))[0] + ".docx"
    if not args.date:
        import datetime
        args.date = datetime.date.today().strftime("%Y.%m.%d")

    if os.path.exists(out):
        print("NOTE: overwriting existing file -> %s" % out)

    convert(args.source, out, args)
    print("Saved DOCX: %s" % out)

    if args.pdf:
        pdf = export_pdf(out)
        if pdf:
            print("Saved PDF:  %s" % pdf)


if __name__ == "__main__":
    main()
